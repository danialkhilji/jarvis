from pathlib import Path
import config

SKILL_NAME = "edit_word"
SKILL_DESCRIPTION = (
    "Read or edit Word documents (.docx). "
    "Actions: read, append_paragraph, replace_text, add_heading. "
    "Returns document text or confirmation of changes."
)


def run(
    file_path: str,
    action: str,
    output_path: str = "",
    search_text: str = "",
    replace_text: str = "",
    paragraph_text: str = "",
    heading_text: str = "",
    heading_level: int = 1,
) -> str:
    """
    Work with Word documents.

    Args:
        file_path: Absolute path to the .docx file.
        action: One of 'read', 'append_paragraph', 'replace_text', 'add_heading'.
        output_path: For write actions, where to save. Defaults to overwriting original.
        search_text: Text to find (for 'replace_text').
        replace_text: Replacement text (for 'replace_text').
        paragraph_text: New paragraph content (for 'append_paragraph').
        heading_text: Heading content (for 'add_heading').
        heading_level: Heading level 1-9 (for 'add_heading').
    """
    try:
        import docx
    except ImportError:
        return "Error: python-docx is not installed. Run: pip install python-docx"

    src = Path(file_path)
    if not src.exists():
        return f"Error: File not found: {file_path}"
    if src.suffix.lower() != ".docx":
        return "Error: Only .docx files are supported."

    allowed = any(str(src).startswith(d) for d in config.ALLOWED_DIRECTORIES)
    if not allowed:
        return "Error: Access denied. Path must be under an allowed directory."

    doc = docx.Document(str(src))

    if action == "read":
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text[:5000] + ("\n\n[truncated...]" if len(text) > 5000 else "")

    dest = Path(output_path) if output_path else src

    if action == "append_paragraph":
        if not paragraph_text:
            return "Error: 'append_paragraph' requires paragraph_text."
        doc.add_paragraph(paragraph_text)

    elif action == "replace_text":
        if not search_text:
            return "Error: 'replace_text' requires search_text."
        count = 0
        for para in doc.paragraphs:
            if search_text in para.text:
                for run in para.runs:
                    if search_text in run.text:
                        run.text = run.text.replace(search_text, replace_text)
                        count += 1
        if count == 0:
            return f"Text '{search_text}' not found in document."

    elif action == "add_heading":
        if not heading_text:
            return "Error: 'add_heading' requires heading_text."
        doc.add_heading(heading_text, level=heading_level)

    else:
        return f"Error: Unknown action '{action}'. Choose from: read, append_paragraph, replace_text, add_heading"

    doc.save(str(dest))
    return f"Word document saved: {dest}"
